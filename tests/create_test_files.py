from docx import Document

def create_sample_docx(filename="sample_resume.docx"):
    doc = Document()
    doc.add_heading('John Doe', 0)
    doc.add_paragraph('Software Engineer with 5 years of experience in Python and FastAPI.')
    
    doc.add_heading('Experience', level=1)
    doc.add_paragraph('Senior Developer at Tech Corp')
    doc.add_paragraph(' - Built REST APIs')
    doc.add_paragraph(' - Optimized database queries')
    
    doc.add_heading('Education', level=1)
    doc.add_paragraph('B.S. Computer Science, University of Technology')
    
    doc.add_heading('Skills', level=1)
    doc.add_paragraph('Python, FastAPI, Docker, SQL')

    doc.save(filename)
    print(f"Created {filename}")

if __name__ == "__main__":
    create_sample_docx()
